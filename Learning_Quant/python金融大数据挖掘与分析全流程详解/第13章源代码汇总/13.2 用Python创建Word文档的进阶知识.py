# =============================================================================
# 13.2 Python操作Word进阶
# =============================================================================

import docx  # 下面的import的代码其实都可以写到这个import下面
file = docx.Document()

# # 1.设置中文字体
from docx.oxml.ns import qn
file.styles['Normal'].font.name = u'微软雅黑'  # 可换成word里面任意字体
file.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 这边记得也得填一下字体名称

# 写入若干段落
# # 2.设置字体颜色和大小
p = file.add_paragraph()
run = p.add_run('螃蟹在剥我的壳，笔记本在写我')
font = run.font
from docx.shared import Pt
font.size = Pt(26)
from docx.shared import RGBColor
font.color.rgb = RGBColor(54, 95, 145)

# # 3.设置字体粗体、斜体和下划线
p = file.add_paragraph()
run = p.add_run('漫天的我落在枫叶上雪花上')
font = run.font
font.bold = True  # 粗体
font.italic = True  # 斜体
font.underline = True  # 下划线

# # 4.设置居中对齐
from docx.enum.text import WD_ALIGN_PARAGRAPH
p = file.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.add_run('而你在想我')

# #  5.设置首行缩进
from docx.shared import Inches
p = file.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.32)
p.add_run('设置首行缩进示例文字')

# # 6.设置行距
from docx.shared import Pt
p = file.add_paragraph()
p.paragraph_format.line_spacing = Pt(16)  # 行距，16磅对应三号字体大小
p.add_run('设置行距示例文字')

# #7.设置段前距和段后距
from docx.shared import Pt
p = file.add_paragraph()
p.paragraph_format.space_before = Pt(14)  # 段前距,14磅对应4号字体大小
p.paragraph_format.space_after = Pt(14)  # 段后距
p.add_run('设置段前段后距示例文字')

# # 8.设置段落序号
file.add_paragraph('点序号', style='List Bullet')
file.add_paragraph('数字序号', style='List Number')

# # 9.设置表格
table = file.add_table(rows=2, cols=3, style='Light Shading Accent 1')
table.cell(0, 0).text = '第一句'  # 第一行第一列
table.cell(0, 1).text = '第二句'  # 第一行第二列
table.cell(0, 2).text = '第三句'  # 第一行第三列
table.cell(1, 0).text = '克制'  # 第二行第一列
table.cell(1, 1).text = '再克制'  # 第二行第二列
table.cell(1, 2).text = '"在吗"'  # 第三行第三列

# # 10.设置图表
file.add_picture('水墨.png', width=Inches(3), height=Inches(3))
last_paragraph = file.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存，得首先创建出保存文件夹，且Word文件不要打开
file.save("三行情书.docx")
print('Word生成完毕！')
