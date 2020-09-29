# =============================================================================
# 13.1-2 Python操作Word基础
# =============================================================================

# 1.创建Word对象
import docx
file = docx.Document()
# 2.添加标题
file.add_heading('三行情书2', level=0)
# 3.添加段落文字
file.add_paragraph('我喜欢你')
file.add_paragraph('上一句话是假的')
file.add_paragraph('上一句话也是假的')
# 4.添加图片
file.add_picture('水墨.png')  # 需要你自己设置一个图片地址
# 5.添加分页符
file.add_page_break()
# 6.添加表格
table = file.add_table(rows=1, cols=3)
table.cell(0,0).text = '克制'
table.cell(0,1).text = '再克制'
table.cell(0,2).text = '在吗'
# 7.文档保存，存储文件夹需提前创建
file.save('E:\\三行情书2.docx')
print('三行情书2生成完毕')
